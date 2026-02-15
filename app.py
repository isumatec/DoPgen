from flask import Flask, render_template, request, send_file
import os
import pandas as pd
from docx import Document
from zipfile import ZipFile
import uuid

app = Flask(__name__)
UPLOAD_FOLDER = "temp"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def cargar_traducciones(excel_file):
    xls = pd.read_excel(excel_file, sheet_name=None)
    df = pd.concat(xls.values(), ignore_index=True)
    df.dropna(subset=["Español"], inplace=True)
    return df

def traducir_documento(ruta_doc, df, idioma, nombre_producto):
    doc = Document(ruta_doc)

    for para in doc.paragraphs:
        for run in para.runs:
            texto = run.text.strip()
            fila = df[df["Español"].str.strip() == texto]
            if not fila.empty:
                run.text = str(fila.iloc[0][idioma]).strip()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        texto = run.text.strip()
                        fila = df[df["Español"].str.strip() == texto]
                        if not fila.empty:
                            run.text = str(fila.iloc[0][idioma]).strip()

    output_name = f"DOP_{nombre_producto}_{idioma}.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_name)
    doc.save(output_path)
    return output_path

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        word_file = request.files["word"]
        excel_file = request.files["excel"]
        idioma = request.form["idioma"]
        nombre_producto = request.form["producto"].replace(" ", "_")

        session_id = str(uuid.uuid4())
        session_folder = os.path.join(UPLOAD_FOLDER, session_id)
        os.makedirs(session_folder, exist_ok=True)

        word_path = os.path.join(session_folder, word_file.filename)
        excel_path = os.path.join(session_folder, excel_file.filename)

        word_file.save(word_path)
        excel_file.save(excel_path)

        df = cargar_traducciones(excel_path)

        idiomas = ["Inglés", "Francés", "Portugués"] if idioma == "Todos" else [idioma]

        archivos_generados = []

        for lang in idiomas:
            path = traducir_documento(word_path, df, lang, nombre_producto)
            archivos_generados.append(path)

        zip_path = os.path.join(session_folder, "DoPgen_Traducciones.zip")
        with ZipFile(zip_path, "w") as zipf:
            for file in archivos_generados:
                zipf.write(file, os.path.basename(file))

        return send_file(zip_path, as_attachment=True)

    return render_template("index.html")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
